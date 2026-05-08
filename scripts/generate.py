#!/usr/bin/env python3
"""FIAP branded document generator. Accepts JSON content, outputs DOCX + PDF."""

import json
import os
import sys

# ---------------------------------------------------------------------------
# Resolve paths relative to this script (works regardless of cwd)
# ---------------------------------------------------------------------------
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SKILL_DIR = os.path.dirname(SCRIPT_DIR)
ASSETS_DIR = os.path.join(SKILL_DIR, "assets")
FONT_DIR = os.path.join(ASSETS_DIR, "montserrat")
HEADER_IMG = os.path.join(ASSETS_DIR, "header-logo.png")

# FIAP brand colors
MAGENTA = (237, 17, 101)
BLACK = (0, 0, 0)
DARK_GRAY = (51, 51, 51)
WHITE = (255, 255, 255)
LIGHT_GRAY = (245, 245, 245)
MID_GRAY = (180, 180, 180)
BORDER_GRAY = (220, 220, 220)

PAGE_W = 210  # A4 width in mm
MARGIN_L = 10
MARGIN_R = 10
CONTENT_W = PAGE_W - MARGIN_L - MARGIN_R  # 190mm usable


def _logo_dimensions(img_path, max_h=14):
    """Return (w, h) in mm preserving aspect ratio, capped by max_h."""
    try:
        from PIL import Image
        img = Image.open(img_path)
        px_w, px_h = img.size
        ratio = px_w / px_h
        h = max_h
        w = h * ratio
        return w, h
    except Exception:
        return 30, 14  # fallback


# ===================================================================
# PDF generator (fpdf2)
# ===================================================================
class FIAPPDF:
    def __init__(self):
        from fpdf import FPDF

        self.pdf = FPDF("P", "mm", "A4")
        self.pdf.set_auto_page_break(auto=True, margin=25)
        self.pdf.add_font("Montserrat", "", os.path.join(FONT_DIR, "Montserrat-Regular.ttf"))
        self.pdf.add_font("Montserrat", "B", os.path.join(FONT_DIR, "Montserrat-Bold.ttf"))
        self.pdf.add_font("Montserrat", "I", os.path.join(FONT_DIR, "Montserrat-Italic.ttf"))
        self.pdf.add_font("Montserrat", "BI", os.path.join(FONT_DIR, "Montserrat-BoldItalic.ttf"))

        _pdf = self.pdf

        # Compute logo size once
        logo_w, logo_h = (30, 14)  # defaults
        if os.path.exists(HEADER_IMG):
            logo_w, logo_h = _logo_dimensions(HEADER_IMG, max_h=14)
        final_logo_w = logo_w
        final_logo_h = logo_h

        def _header():
            if os.path.exists(HEADER_IMG):
                _pdf.image(HEADER_IMG, x=MARGIN_L + 9, y=12, w=final_logo_w, h=final_logo_h)
            _pdf.set_y(12 + final_logo_h + 6)

        def _footer():
            _pdf.set_y(-10)
            _pdf.set_font("Montserrat", "I", 7)
            _pdf.set_text_color(*MID_GRAY)
            _pdf.cell(0, 5, f"{_pdf.page_no()}", align="R")

        _pdf.header = _header  # type: ignore[assignment]
        _pdf.footer = _footer  # type: ignore[assignment]

    # --- Cover helpers -----------------------------------------------------

    def _cover_title(self, text, font_size=28, color=MAGENTA, line_h=14):
        """Render a cover title with word wrapping to respect margins."""
        p = self.pdf
        p.set_font("Montserrat", "B", font_size)
        p.set_text_color(*color)

        # Use multi_cell for wrapping within margins
        p.set_x(MARGIN_L)
        p.multi_cell(CONTENT_W, line_h, text, align="C")

    # --- Section helpers ---------------------------------------------------

    def _section_title(self, text):
        p = self.pdf
        if p.get_y() > 220:
            p.add_page()
        p.ln(6)
        p.set_font("Montserrat", "B", 13)
        p.set_text_color(*MAGENTA)
        p.cell(0, 9, text, new_x="LMARGIN", new_y="NEXT")
        p.set_draw_color(*MAGENTA)
        p.set_line_width(0.8)
        p.line(MARGIN_L, p.get_y(), PAGE_W - MARGIN_R, p.get_y())
        p.ln(4)

    def _subsection_title(self, text):
        p = self.pdf
        p.ln(3)
        p.set_font("Montserrat", "B", 10)
        p.set_text_color(*BLACK)
        p.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")
        p.ln(2)

    def _body(self, text, bold=False, italic=False):
        p = self.pdf
        style = ""
        if bold and italic:
            style = "BI"
        elif bold:
            style = "B"
        elif italic:
            style = "I"
        p.set_font("Montserrat", style, 9)
        p.set_text_color(*DARK_GRAY)
        p.multi_cell(0, 5, text)
        p.ln(2)

    def _bullet(self, text, level=0):
        p = self.pdf
        indent = 8 + level * 6
        marker = "•" if level == 0 else "–"
        p.set_font("Montserrat", "", 9)
        p.set_text_color(*DARK_GRAY)
        x = p.get_x()
        p.set_x(x + indent)
        p.multi_cell(CONTENT_W - indent, 5, f"{marker}  {text}")
        p.ln(1)

    def _add_table(self, headers, rows, col_widths=None):
        p = self.pdf
        total = CONTENT_W
        if col_widths is None:
            col_widths = [total // len(headers)] * len(headers)
            col_widths[-1] = total - sum(col_widths[:-1])

        HEADER_H = 8
        MAX_Y = 265

        p.set_auto_page_break(auto=False)

        def _draw_header():
            p.set_font("Montserrat", "B", 8)
            p.set_fill_color(*MAGENTA)
            p.set_text_color(*WHITE)
            if p.get_y() + HEADER_H + 9 > MAX_Y:
                p.add_page()
            y0 = p.get_y()
            for i, h in enumerate(headers):
                p.set_xy(MARGIN_L + sum(col_widths[:i]), y0)
                p.cell(col_widths[i], HEADER_H, h, border=0, fill=True, align="C")
            p.set_y(y0 + HEADER_H)

        _draw_header()
        p.set_font("Montserrat", "", 8)
        p.set_draw_color(*BORDER_GRAY)
        p.set_line_width(0.2)

        for row_idx, row in enumerate(rows):
            bg = LIGHT_GRAY if row_idx % 2 == 1 else WHITE
            p.set_fill_color(*bg)
            p.set_text_color(*DARK_GRAY)

            max_lines = 1
            for i, cell in enumerate(row):
                lines = p.multi_cell(col_widths[i] - 4, 5, cell, dry_run=True, output="LINES")
                max_lines = max(max_lines, len(lines))
            row_h = max(7, max_lines * 5 + 2)

            if p.get_y() + row_h > MAX_Y:
                p.add_page()
                _draw_header()
                p.set_font("Montserrat", "", 8)
                p.set_draw_color(*BORDER_GRAY)
                p.set_line_width(0.2)
                bg = LIGHT_GRAY if row_idx % 2 == 1 else WHITE
                p.set_fill_color(*bg)
                p.set_text_color(*DARK_GRAY)

            y_row = p.get_y()
            p.rect(MARGIN_L, y_row, sum(col_widths), row_h, style="F")
            for i in range(len(headers)):
                p.rect(MARGIN_L + sum(col_widths[:i]), y_row, col_widths[i], row_h, style="D")
            for i, cell in enumerate(row):
                p.set_xy(MARGIN_L + sum(col_widths[:i]) + 2, y_row + 1)
                p.multi_cell(col_widths[i] - 4, 5, cell)

            p.set_y(y_row + row_h)

        p.set_draw_color(*MAGENTA)
        p.set_line_width(0.5)
        p.line(MARGIN_L, p.get_y(), MARGIN_L + sum(col_widths), p.get_y())
        p.ln(4)

        p.set_auto_page_break(auto=True, margin=25)

    # --- Recursive section renderer ----------------------------------------

    def _render_section(self, section):
        self._section_title(section["title"])
        self._render_content(section)

    def _render_subsection(self, sub):
        self._subsection_title(sub["title"])
        self._render_content(sub)

    def _render_content(self, obj):
        bodies = obj.get("body", [])
        for i, text in enumerate(bodies):
            bold = obj.get("bold_body", False) and i == 0
            self._body(text, bold=bold)

        for b in obj.get("bullets", []):
            if isinstance(b, dict):
                self._bullet(b["text"], b.get("level", 0))
            else:
                self._bullet(b, 0)

        for t in obj.get("tables", []):
            self._add_table(t["headers"], t["rows"], t.get("col_widths"))

        for sub in obj.get("subsections", []):
            self._render_subsection(sub)

    # --- Build full document ------------------------------------------------

    def build(self, content: dict) -> "fpdf.FPDF":
        p = self.pdf

        # Cover page
        p.add_page()
        p.set_y(45)

        # Title with wrapping
        self._cover_title(content["title"], font_size=28, color=MAGENTA, line_h=14)
        p.ln(2)

        for line in content.get("subtitle", []):
            p.set_font("Montserrat", "B", 22)
            p.set_text_color(*BLACK)
            p.set_x(MARGIN_L)
            p.multi_cell(CONTENT_W, 12, line, align="C")
        p.ln(8)

        note = content.get("subtitle_note")
        if note:
            p.set_font("Montserrat", "I", 12)
            p.set_text_color(*DARK_GRAY)
            p.set_x(MARGIN_L)
            p.multi_cell(CONTENT_W, 8, note, align="C")
        p.ln(16)

        for m in content.get("meta", []):
            p.set_font("Montserrat", "B", 9)
            p.cell(45, 7, m["label"])
            p.set_font("Montserrat", "", 9)
            p.multi_cell(CONTENT_W - 45, 7, m["value"])
            p.ln(0)

        sections = content.get("sections", [])
        if sections:
            p.add_page()
            self._render_section(sections[0])
        for section in sections[1:]:
            self._render_section(section)

        return p


# ===================================================================
# DOCX generator (python-docx)
# ===================================================================
class FIAPDOCX:
    def __init__(self):
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self.doc = Document()
        self.Pt = Pt
        self.RGBColor = RGBColor
        self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH

    def _section_title(self, text):
        doc = self.doc
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Montserrat"
        run.font.size = self.Pt(13)
        run.font.bold = True
        run.font.color.rgb = self.RGBColor(*MAGENTA)
        para.space_after = self.Pt(6)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "ED1165")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _subsection_title(self, text):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Montserrat"
        run.font.size = self.Pt(10)
        run.font.bold = True
        run.font.color.rgb = self.RGBColor(*BLACK)

    def _body(self, text, bold=False):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Montserrat"
        run.font.size = self.Pt(9)
        run.font.bold = bold
        run.font.color.rgb = self.RGBColor(*DARK_GRAY)

    def _bullet(self, text, level=0):
        para = self.doc.add_paragraph(style="List Bullet")
        if level == 1:
            para.paragraph_format.left_indent = self.Pt(36)
        run = para.add_run(text)
        run.font.name = "Montserrat"
        run.font.size = self.Pt(9)
        run.font.color.rgb = self.RGBColor(*DARK_GRAY)

    def _add_table(self, headers, rows):
        doc = self.doc
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"

        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.font.name = "Montserrat"
            run.font.size = self.Pt(8)
            run.font.bold = True
            run.font.color.rgb = self.RGBColor(*WHITE)
            cell.paragraphs[0].alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "ED1165")
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)

        for row_idx, row in enumerate(rows):
            for i, val in enumerate(row):
                cell = table.rows[row_idx + 1].cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(val)
                run.font.name = "Montserrat"
                run.font.size = self.Pt(8)
                run.font.color.rgb = self.RGBColor(*DARK_GRAY)
                if row_idx % 2 == 1:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "F5F5F5")
                    shading.set(qn("w:val"), "clear")
                    cell._tc.get_or_add_tcPr().append(shading)

        doc.add_paragraph()

    def _render_content(self, obj):
        bodies = obj.get("body", [])
        for i, text in enumerate(bodies):
            bold = obj.get("bold_body", False) and i == 0
            self._body(text, bold=bold)

        for b in obj.get("bullets", []):
            if isinstance(b, dict):
                self._bullet(b["text"], b.get("level", 0))
            else:
                self._bullet(b, 0)

        for t in obj.get("tables", []):
            self._add_table(t["headers"], t["rows"])

        for sub in obj.get("subsections", []):
            self._subsection_title(sub["title"])
            self._render_content(sub)

    def build(self, content: dict):
        doc = self.doc

        # Cover page
        para = doc.add_paragraph()
        para.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(content["title"])
        run.font.name = "Montserrat"
        run.font.size = self.Pt(28)
        run.font.bold = True
        run.font.color.rgb = self.RGBColor(*MAGENTA)

        for line in content.get("subtitle", []):
            para = doc.add_paragraph()
            para.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line)
            run.font.name = "Montserrat"
            run.font.size = self.Pt(22)
            run.font.bold = True
            run.font.color.rgb = self.RGBColor(*BLACK)

        note = content.get("subtitle_note")
        if note:
            para = doc.add_paragraph()
            para.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(note)
            run.font.name = "Montserrat"
            run.font.size = self.Pt(12)
            run.font.italic = True
            run.font.color.rgb = self.RGBColor(*DARK_GRAY)

        doc.add_paragraph()

        for m in content.get("meta", []):
            para = doc.add_paragraph()
            run_label = para.add_run(m["label"] + " ")
            run_label.font.name = "Montserrat"
            run_label.font.size = self.Pt(9)
            run_label.font.bold = True
            # Handle multiline values (integrants list with \n)
            lines = m["value"].split("\n")
            for li, line in enumerate(lines):
                if li > 0:
                    run_value.add_break()
                    run_value = para.add_run(line)
                    run_value.font.name = "Montserrat"
                else:
                    run_value = para.add_run(line)
                    run_value.font.name = "Montserrat"
                run_value.font.size = self.Pt(9)

        for i, section in enumerate(content.get("sections", [])):
            if i == 0:
                doc.add_page_break()
            else:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = self.Pt(12)
            self._section_title(section["title"])
            self._render_content(section)

        return doc


# ===================================================================
# CLI entry point
# ===================================================================
def main():
    if len(sys.argv) < 2:
        print("Usage: python generate.py <content.json> [--pdf-only|--docx-only]")
        sys.exit(1)

    content_path = sys.argv[1]
    with open(content_path, encoding="utf-8") as f:
        content = json.load(f)

    flags = set(sys.argv[2:])
    do_pdf = "--docx-only" not in flags
    do_docx = "--pdf-only" not in flags

    output_dir = content.get("output_dir") or os.path.join(os.path.expanduser("~"), "Downloads")
    os.makedirs(output_dir, exist_ok=True)
    base_name = content.get("filename", "FIAP_Document")

    results = []

    if do_pdf:
        pdf_builder = FIAPPDF()
        pdf = pdf_builder.build(content)
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        pdf.output(pdf_path)
        results.append(f"PDF: {pdf_path}")

    if do_docx:
        docx_builder = FIAPDOCX()
        doc = docx_builder.build(content)
        docx_path = os.path.join(output_dir, f"{base_name}.docx")
        doc.save(docx_path)
        results.append(f"DOCX: {docx_path}")

    for r in results:
        print(r)


if __name__ == "__main__":
    main()
